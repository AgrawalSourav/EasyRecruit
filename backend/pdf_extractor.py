"""
PDF and DOCX text extraction utilities
Handles various document formats with fallback methods
"""

import io
import logging
from typing import Optional

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file with multiple fallback methods
        """
        text = ""
        
        # Method 1: Try pdfplumber first (better for complex layouts)
        if pdfplumber:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    logger.info("Successfully extracted text using pdfplumber")
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 2: Fallback to PyPDF2
        if PyPDF2:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    logger.info("Successfully extracted text using PyPDF2")
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # If both methods fail
        logger.error("Failed to extract text from PDF using all available methods")
        return "Error: Could not extract text from PDF. Please ensure the PDF contains readable text."
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        """
        if not Document:
            return "Error: python-docx not installed. Cannot process DOCX files."
        
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info("Successfully extracted text from DOCX")
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"Error: Could not extract text from DOCX - {str(e)}"
    
    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        """
        Extract text based on file extension
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return DocumentExtractor.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return DocumentExtractor.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.txt'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except:
                    return "Error: Could not decode text file"
        else:
            return f"Error: Unsupported file format. Supported formats: PDF, DOCX, TXT"
